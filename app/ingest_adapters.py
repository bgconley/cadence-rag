from __future__ import annotations

import csv
import html
import json
import logging
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Mapping, Optional, Sequence, Tuple

from .config import settings
from .schemas import TranscriptPayload

TURN_LIST_KEYS = (
    "content",
    "utterances",
    "turns",
    "segments",
    "items",
    "messages",
)
TEXT_KEYS = ("text", "transcript", "utterance", "content", "message", "body")
SPEAKER_KEYS = (
    "speaker",
    "speaker_name",
    "speakerName",
    "participant",
    "participant_name",
    "role",
    "user",
    "name",
)
START_KEYS = (
    "start_ts_ms",
    "start_ms",
    "start",
    "start_time",
    "startTime",
    "start_sec",
    "start_seconds",
)
END_KEYS = (
    "end_ts_ms",
    "end_ms",
    "end",
    "end_time",
    "endTime",
    "end_sec",
    "end_seconds",
)
DURATION_KEYS = ("duration_ms", "duration", "duration_s")
HTML_TAG_RE = re.compile(r"<[^>]+>")
TIMESTAMP_RE = re.compile(
    r"^(?:(?P<hours>\d{1,2}):)?(?P<minutes>[0-5]?\d):(?P<seconds>[0-5]?\d(?:\.\d+)?)$"
)
SPEAKER_PREFIX_RE = re.compile(r"^\s*(?P<speaker>[^:]{1,80}):\s*(?P<text>.+?)\s*$")
MARKDOWN_SPEAKER_RE = re.compile(
    r"^\s*\*\*(?P<speaker>.+?)\*\*\s*:\s*(?P<text>.*?)\s*$"
)
MARKDOWN_TS_RE = re.compile(
    r"^\s*\*?(?:(?P<hours>\d{1,2}):)?(?P<minutes>[0-5]?\d):(?P<seconds>[0-5]?\d)\*?\s*$"
)
logger = logging.getLogger(__name__)


def load_transcript_payload(path: Path, *, format_hint: str = "json_turns") -> TranscriptPayload:
    normalized_hint = (format_hint or "json_turns").lower()
    if normalized_hint == "markdown_turns":
        normalized = _normalize_markdown_transcript(path.read_text(encoding="utf-8"))
    else:
        normalized = _load_and_normalize_json_transcript(path, format_hint=normalized_hint)
    return TranscriptPayload.model_validate({"format": "json_turns", "content": normalized})


def _load_and_normalize_json_transcript(path: Path, *, format_hint: str) -> list[dict[str, Any]]:
    if format_hint not in {"json_turns", "auto"}:
        raise ValueError(f"unsupported transcript format hint: {format_hint}")

    raw_text = path.read_text(encoding="utf-8", errors="replace")
    try:
        raw_obj = json.loads(raw_text)
        return _normalize_transcript_object(raw_obj, format_hint=format_hint)
    except Exception:
        if format_hint != "auto":
            raise
        return _normalize_markdown_transcript(raw_text)


def load_analysis_content(path: Path, *, format_hint: str = "auto") -> str:
    normalized_hint = (format_hint or "auto").lower()
    if normalized_hint == "auto":
        suffix = path.suffix.lower()
        if suffix in {".md", ".markdown", ".txt", ".log"}:
            normalized_hint = "text"
        elif suffix == ".csv":
            normalized_hint = "csv"
        elif suffix == ".tsv":
            normalized_hint = "tsv"
        elif suffix == ".json":
            normalized_hint = "json"
        elif suffix in {".html", ".htm"}:
            normalized_hint = "html"
        elif suffix == ".docx":
            normalized_hint = "docx"
        elif suffix == ".pdf":
            normalized_hint = "pdf"
        else:
            normalized_hint = "text"

    if normalized_hint in {"text", "markdown"}:
        return path.read_text(encoding="utf-8", errors="replace")
    if normalized_hint == "csv":
        return _csv_to_markdown(path, delimiter=",")
    if normalized_hint == "tsv":
        return _csv_to_markdown(path, delimiter="\t")
    if normalized_hint == "json":
        return _json_to_text(path)
    if normalized_hint == "html":
        raw = path.read_text(encoding="utf-8", errors="replace")
        return _html_to_text(raw)
    if normalized_hint == "docx":
        return _docx_to_text(path)
    if normalized_hint == "pdf":
        return _pdf_to_text(path)
    raise ValueError(f"unsupported analysis format: {format_hint}")


def _docx_to_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ValueError(
            "DOCX support requires python-docx; add it to project dependencies"
        ) from exc

    document = Document(path)
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)
        if not rows:
            continue
        header = rows[0]
        body = rows[1:]
        lines.append("")
        lines.append("| " + " | ".join(_md_escape(cell) for cell in header) + " |")
        lines.append("| " + " | ".join("---" for _ in header) + " |")
        for row in body:
            padded = row + [""] * max(0, len(header) - len(row))
            cells = padded[: len(header)]
            lines.append("| " + " | ".join(_md_escape(cell) for cell in cells) + " |")

    return "\n".join(lines).strip()


def _pdf_to_text(path: Path) -> str:
    extracted_text, page_count = _extract_pdf_text_with_pypdf(path)

    if not settings.analysis_pdf_ocr_enabled:
        return extracted_text

    if not _should_run_pdf_ocr(extracted_text, page_count):
        return extracted_text

    ocr_text = _pdf_to_text_via_ocrmypdf(path)
    if not ocr_text:
        return extracted_text

    if _is_better_pdf_text(candidate=ocr_text, baseline=extracted_text):
        return ocr_text
    return extracted_text or ocr_text


def _extract_pdf_text_with_pypdf(path: Path) -> tuple[str, int]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ValueError(
            "PDF support requires pypdf; add it to project dependencies"
        ) from exc

    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        pages.append(f"## Page {index}\n\n{text}")
    return "\n\n".join(pages).strip(), len(reader.pages)


def _should_run_pdf_ocr(extracted_text: str, page_count: int) -> bool:
    if page_count <= 0:
        return False
    if page_count > settings.analysis_pdf_ocr_max_pages:
        logger.info(
            "pdf_ocr.skipped page_count=%s max_pages=%s",
            page_count,
            settings.analysis_pdf_ocr_max_pages,
        )
        return False
    if settings.analysis_pdf_ocr_force:
        return True

    text_char_count, alpha_ratio = _pdf_text_quality(extracted_text)
    if text_char_count < settings.analysis_pdf_ocr_min_chars:
        return True
    if alpha_ratio < settings.analysis_pdf_ocr_min_alpha_ratio:
        return True
    return False


def _pdf_to_text_via_ocrmypdf(path: Path) -> str:
    ocr_command = settings.analysis_pdf_ocr_command.strip()
    if not ocr_command:
        return ""

    if shutil.which(ocr_command) is None:
        logger.warning("pdf_ocr.command_missing command=%s", ocr_command)
        return ""

    with tempfile.TemporaryDirectory(prefix="cadence-rag-ocr-") as temp_dir:
        sidecar_path = Path(temp_dir) / "ocr_sidecar.txt"
        output_pdf_path = Path(temp_dir) / "ocr_output.pdf"
        cmd = [
            ocr_command,
            "--skip-text",
            "--quiet",
            "--sidecar",
            str(sidecar_path),
            "--language",
            settings.analysis_pdf_ocr_languages,
            str(path),
            str(output_pdf_path),
        ]
        try:
            completed = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=settings.analysis_pdf_ocr_timeout_s,
                check=False,
            )
        except Exception as exc:
            logger.warning("pdf_ocr.execution_error error=%s", str(exc))
            return ""

        if completed.returncode != 0:
            stderr = (completed.stderr or "").strip()
            logger.warning(
                "pdf_ocr.failed command=%s returncode=%s stderr=%s",
                ocr_command,
                completed.returncode,
                stderr[:200],
            )
            return ""

        if sidecar_path.exists():
            return sidecar_path.read_text(encoding="utf-8", errors="replace").strip()

        text, _ = _extract_pdf_text_with_pypdf(output_pdf_path)
        return text


def _pdf_text_quality(text: str) -> tuple[int, float]:
    non_ws_count = sum(1 for ch in text if not ch.isspace())
    if non_ws_count == 0:
        return 0, 0.0
    alpha_chars = sum(1 for ch in text if ch.isalpha())
    return non_ws_count, alpha_chars / non_ws_count


def _is_better_pdf_text(*, candidate: str, baseline: str) -> bool:
    candidate_chars, candidate_alpha_ratio = _pdf_text_quality(candidate)
    baseline_chars, baseline_alpha_ratio = _pdf_text_quality(baseline)
    if candidate_chars == 0:
        return False
    if baseline_chars == 0:
        return True
    if candidate_chars >= baseline_chars * 1.15:
        return True
    return candidate_alpha_ratio > baseline_alpha_ratio + 0.05


def _normalize_transcript_object(raw: Any, *, format_hint: str) -> list[dict[str, Any]]:
    if format_hint not in {"json_turns", "auto"}:
        raise ValueError(f"unsupported transcript format hint: {format_hint}")

    items = _extract_turn_items(raw)
    result: list[dict[str, Any]] = []
    cursor_ms = 0
    for item in items:
        parsed = _normalize_turn(item, cursor_ms=cursor_ms)
        if parsed is None:
            continue
        result.append(parsed)
        cursor_ms = parsed["end_ts_ms"]
    if not result:
        raise ValueError("transcript contains no usable utterances")
    return result


def _normalize_markdown_transcript(raw_text: str) -> list[dict[str, Any]]:
    lines = [line.rstrip("\n") for line in raw_text.splitlines()]
    entries: list[dict[str, Any]] = []
    current: Optional[dict[str, Any]] = None

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue

        speaker_match = MARKDOWN_SPEAKER_RE.match(line)
        if speaker_match:
            if current is not None:
                entries.append(current)
            current = {
                "speaker": speaker_match.group("speaker").strip() or "UNKNOWN",
                "text": speaker_match.group("text").strip(),
                "start_ts_ms": None,
            }
            continue

        timestamp_ms = _parse_markdown_timestamp(line)
        if timestamp_ms is not None:
            if current is None:
                current = {"speaker": "UNKNOWN", "text": "", "start_ts_ms": timestamp_ms}
            else:
                current["start_ts_ms"] = timestamp_ms
            continue

        if current is None:
            speaker, text = _parse_speaker_prefixed_line(line)
            current = {"speaker": speaker, "text": text, "start_ts_ms": None}
        else:
            continuation = line
            if continuation:
                current["text"] = f"{current['text']} {continuation}".strip()

    if current is not None:
        entries.append(current)

    normalized: list[dict[str, Any]] = []
    cursor_ms = 0
    for idx, item in enumerate(entries):
        text_val = str(item.get("text", "")).strip()
        if not text_val:
            continue

        start_ms = item.get("start_ts_ms")
        if start_ms is None:
            start_ms = cursor_ms
        start_ms = int(start_ms)

        next_start_ms = _next_explicit_start(entries, idx + 1)
        if next_start_ms is None:
            end_ms = start_ms + 1000
        else:
            end_ms = max(start_ms + 1, int(next_start_ms))

        normalized.append(
            {
                "speaker": str(item.get("speaker") or "UNKNOWN"),
                "start_ts_ms": start_ms,
                "end_ts_ms": end_ms,
                "text": text_val,
            }
        )
        cursor_ms = end_ms

    if not normalized:
        raise ValueError("transcript contains no usable utterances")
    return normalized


def _parse_markdown_timestamp(line: str) -> Optional[int]:
    match = MARKDOWN_TS_RE.match(line)
    if not match:
        return None
    hours = int(match.group("hours") or 0)
    minutes = int(match.group("minutes"))
    seconds = int(match.group("seconds"))
    total_ms = (hours * 3600 + minutes * 60 + seconds) * 1000
    return total_ms


def _next_explicit_start(entries: Sequence[Mapping[str, Any]], start_idx: int) -> Optional[int]:
    for i in range(start_idx, len(entries)):
        value = entries[i].get("start_ts_ms")
        if value is not None:
            return int(value)
    return None


def _extract_turn_items(raw: Any) -> list[Any]:
    if isinstance(raw, list):
        return raw
    if isinstance(raw, Mapping):
        content = raw.get("content")
        if isinstance(content, list):
            return content
        for key in TURN_LIST_KEYS:
            value = raw.get(key)
            if isinstance(value, list):
                return value
        nested_results = raw.get("results")
        if isinstance(nested_results, Mapping):
            for key in TURN_LIST_KEYS:
                value = nested_results.get(key)
                if isinstance(value, list):
                    return value
        # Single utterance-like object fallback.
        if any(key in raw for key in TEXT_KEYS):
            return [raw]
    raise ValueError("unsupported transcript JSON shape")


def _normalize_turn(item: Any, *, cursor_ms: int) -> Optional[dict[str, Any]]:
    if isinstance(item, str):
        speaker, text = _parse_speaker_prefixed_line(item)
        if not text:
            return None
        start_ms = cursor_ms
        end_ms = start_ms + 1000
        return {
            "speaker": speaker,
            "start_ts_ms": start_ms,
            "end_ts_ms": end_ms,
            "text": text,
        }

    if not isinstance(item, Mapping):
        return None

    text_val = _extract_text(item)
    if not text_val:
        return None

    speaker = _extract_speaker(item) or "UNKNOWN"
    start_key, start_raw = _extract_with_key(item, START_KEYS)
    end_key, end_raw = _extract_with_key(item, END_KEYS)
    duration_key, duration_raw = _extract_with_key(item, DURATION_KEYS)

    start_ms = _to_milliseconds(start_raw, source_key=start_key)
    end_ms = _to_milliseconds(end_raw, source_key=end_key)
    duration_ms = _to_milliseconds(duration_raw, source_key=duration_key)

    if start_ms is None:
        start_ms = cursor_ms
    if end_ms is None and duration_ms is not None:
        end_ms = start_ms + duration_ms
    if end_ms is None:
        end_ms = max(start_ms + 1000, cursor_ms + 1000)
    if end_ms <= start_ms:
        end_ms = start_ms + 1

    return {
        "speaker": speaker,
        "speaker_id": _as_optional_str(item.get("speaker_id")),
        "start_ts_ms": int(start_ms),
        "end_ts_ms": int(end_ms),
        "confidence": _to_float(item.get("confidence")),
        "text": text_val,
    }


def _parse_speaker_prefixed_line(line: str) -> Tuple[str, str]:
    normalized = line.strip()
    if not normalized:
        return "UNKNOWN", ""
    match = SPEAKER_PREFIX_RE.match(normalized)
    if not match:
        return "UNKNOWN", normalized
    speaker = match.group("speaker").strip() or "UNKNOWN"
    text = match.group("text").strip()
    return speaker, text


def _extract_text(item: Mapping[str, Any]) -> str:
    _, value = _extract_with_key(item, TEXT_KEYS)
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, (int, float)):
        return str(value)
    alternatives = item.get("alternatives")
    if isinstance(alternatives, Sequence) and alternatives:
        first = alternatives[0]
        if isinstance(first, Mapping):
            _, alt_text = _extract_with_key(first, TEXT_KEYS)
            if isinstance(alt_text, str):
                return alt_text.strip()
    words = item.get("words")
    if isinstance(words, Sequence):
        tokens = []
        for word in words:
            if isinstance(word, Mapping):
                token = word.get("word") or word.get("text")
                if isinstance(token, str):
                    stripped = token.strip()
                    if stripped:
                        tokens.append(stripped)
        if tokens:
            return " ".join(tokens)
    return ""


def _extract_speaker(item: Mapping[str, Any]) -> str:
    _, value = _extract_with_key(item, SPEAKER_KEYS)
    return _as_optional_str(value) or ""


def _extract_with_key(item: Mapping[str, Any], keys: Sequence[str]) -> Tuple[str, Any]:
    for key in keys:
        if key in item:
            value = item[key]
            if value is not None:
                return key, value
    return "", None


def _to_float(value: Any) -> Optional[float]:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        stripped = value.strip()
        if not stripped:
            return None
        try:
            return float(stripped)
        except ValueError:
            return None
    return None


def _to_milliseconds(value: Any, *, source_key: str) -> Optional[int]:
    if isinstance(value, str):
        stripped = value.strip()
        if not stripped:
            return None
        parsed_hms = _parse_hms_to_seconds(stripped)
        if parsed_hms is not None:
            return int(parsed_hms * 1000)

    numeric = _to_float(value)
    if numeric is None:
        return None

    key = source_key.lower()
    if "ms" in key:
        return int(numeric)
    if "sec" in key:
        return int(numeric * 1000)

    # Heuristic for ambiguous keys (`start`, `end`, `start_time`):
    # - epoch seconds (~1e9) -> seconds
    # - large offsets (>= 10_000) are likely milliseconds
    if numeric >= 1_000_000_000 and numeric < 100_000_000_000:
        return int(numeric * 1000)
    if numeric >= 10_000:
        return int(numeric)
    return int(numeric * 1000)


def _parse_hms_to_seconds(value: str) -> Optional[float]:
    match = TIMESTAMP_RE.match(value)
    if not match:
        return None
    hours = int(match.group("hours") or 0)
    minutes = int(match.group("minutes"))
    seconds = float(match.group("seconds"))
    return float(hours * 3600 + minutes * 60) + seconds


def _as_optional_str(value: Any) -> Optional[str]:
    if value is None:
        return None
    if isinstance(value, str):
        stripped = value.strip()
        return stripped or None
    if isinstance(value, (int, float)):
        return str(value)
    return None


def _csv_to_markdown(path: Path, *, delimiter: str) -> str:
    with path.open("r", encoding="utf-8", errors="replace", newline="") as handle:
        reader = csv.reader(handle, delimiter=delimiter)
        rows = [row for row in reader]

    if not rows:
        return ""

    header = rows[0]
    body = rows[1:]
    lines = [
        "| " + " | ".join(_md_escape(cell) for cell in header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in body:
        padded = row + [""] * max(0, len(header) - len(row))
        cells = padded[: len(header)]
        lines.append("| " + " | ".join(_md_escape(cell) for cell in cells) + " |")
    return "\n".join(lines)


def _json_to_text(path: Path) -> str:
    parsed = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    if isinstance(parsed, list) and parsed and all(isinstance(row, Mapping) for row in parsed):
        return _records_to_markdown_table(parsed)
    if isinstance(parsed, Mapping):
        for key in ("rows", "data", "items", "records"):
            value = parsed.get(key)
            if isinstance(value, list) and value and all(
                isinstance(row, Mapping) for row in value
            ):
                return f"## {key}\n\n{_records_to_markdown_table(value)}"
    return json.dumps(parsed, indent=2, ensure_ascii=False)


def _records_to_markdown_table(records: Sequence[Mapping[str, Any]]) -> str:
    columns: list[str] = []
    seen: set[str] = set()
    for row in records:
        for key in row.keys():
            if key not in seen:
                seen.add(key)
                columns.append(str(key))

    lines = [
        "| " + " | ".join(_md_escape(col) for col in columns) + " |",
        "| " + " | ".join("---" for _ in columns) + " |",
    ]
    for row in records:
        cells = []
        for col in columns:
            value = row.get(col, "")
            if isinstance(value, (dict, list)):
                cells.append(json.dumps(value, ensure_ascii=False))
            else:
                cells.append("" if value is None else str(value))
        lines.append("| " + " | ".join(_md_escape(cell) for cell in cells) + " |")
    return "\n".join(lines)


def _html_to_text(raw_html: str) -> str:
    without_tags = HTML_TAG_RE.sub(" ", raw_html)
    unescaped = html.unescape(without_tags)
    lines = [re.sub(r"\s+", " ", line).strip() for line in unescaped.splitlines()]
    return "\n".join(line for line in lines if line)


def _md_escape(value: Any) -> str:
    text = str(value)
    text = text.replace("|", "\\|")
    text = text.replace("\n", " ").replace("\r", " ")
    return text.strip()
